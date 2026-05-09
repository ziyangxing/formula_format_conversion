#!/usr/bin/env python3
"""
md2docx.py - Markdown → DOCX 转换器

将 AI 对话导出的 .md 文件转换为 .docx，保留公式分隔符 ($, $$, etc.)
供 convert.py 后续转换为专业公式格式。

支持: 标题, 段落, 代码块, 行内代码, 粗体/斜体, 列表, 表格, 引用, 水平线
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class HeadingNumbering:
    """标题自动编号器。
    栈式算法：遇到标题时更新对应层级的计数器，产生如 "1.2.3" 的编号。
    如果标题原文已含手动编号（一、/ 1. / (1) 等），则跳过自动编号。
    """
    # 匹配标题开头已有的手动编号
    MANUAL_NUMBER = re.compile(
        r'^[\s]*('
        r'[一-鿿]{1,3}[、．.]'       # 一、二、三、
        r'|\d+[\.\、\)]\s*'                   # 1. 1) 1、
        r'|[（(]\d+[）)]'                      # (1) （1）
        r'|[ivxIVX]+[\.、]\s*'                # i. ii. iv.
        r')'
    )

    def __init__(self):
        self._counters = {}

    def need_auto_number(self, raw_title):
        """检查标题是否需要自动编号（原标题无手动编号）。"""
        return not self.MANUAL_NUMBER.match(raw_title.strip())

    def get_number(self, level):
        """返回当前标题编号，如 '1.2'。"""
        if level < 1:
            return ''
        for lv in list(self._counters.keys()):
            if lv > level:
                self._counters[lv] = 0
        self._counters[level] = self._counters.get(level, 0) + 1
        path = [str(self._counters[lv]) for lv in sorted(self._counters)
                if lv <= level and self._counters.get(lv, 0) > 0]
        return '.'.join(path)


def md_to_docx(md_path: str, docx_path: str):
    """将 .md 文件转换为 .docx。"""
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()
    _setup_styles(doc)

    numbering = HeadingNumbering()

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            block_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block_lines.append(lines[i])
                i += 1
            _add_code_block(doc, block_lines, lang)
            i += 1  # skip closing ```
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|?[\s\-:|]+\|?$', lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # 空白行
        if not line.strip():
            i += 1
            continue

        # 标题
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            raw_title = header_match.group(2)
            if numbering.need_auto_number(raw_title):
                num = numbering.get_number(level)
                title = f'{num} {raw_title}' if num else raw_title
            else:
                title = raw_title
            _add_heading(doc, title, min(level, 6))
            i += 1
            continue

        # 水平线
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # 引用
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i]))
                i += 1
            _add_blockquote(doc, '\n'.join(quote_lines))
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if list_match:
            list_items = []
            indent = len(list_match.group(1))
            while i < len(lines):
                lm = re.match(r'^(\s*)[-*+]\s+(.+)$', lines[i])
                if lm and len(lm.group(1)) == indent:
                    list_items.append(lm.group(2))
                    i += 1
                elif not lines[i].strip():
                    i += 1
                    break
                else:
                    break
            _add_list(doc, list_items, ordered=False)
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ol_match:
            list_items = []
            indent = len(ol_match.group(1))
            while i < len(lines):
                om = re.match(r'^(\s*)\d+\.\s+(.+)$', lines[i])
                if om and len(om.group(1)) == indent:
                    list_items.append(om.group(2))
                    i += 1
                elif not lines[i].strip():
                    i += 1
                    break
                else:
                    break
            _add_list(doc, list_items, ordered=True)
            continue

        # 普通段落
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('```'):
            para_lines.append(lines[i])
            i += 1
        _add_paragraph(doc, ' '.join(para_lines))

    doc.save(docx_path)
    return docx_path


def _setup_styles(doc):
    """设置默认样式：字体、1.5倍行距、首行缩进。"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)  # ~2个中文字符宽度


def _add_paragraph(doc, text):
    """添加段落，处理行内格式。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    _add_formatted_runs(p, text)


def _add_heading(doc, text, level):
    """添加标题，使用 Word 内置标题样式 (1-6级，视觉区分)。"""
    p = doc.add_heading(level=level)
    _clear_heading_run(p)
    _add_formatted_runs(p, text)

    # 补充字号设定确保层级区分
    font_sizes = {1: Pt(22), 2: Pt(16), 3: Pt(14), 4: Pt(13), 5: Pt(12), 6: Pt(11)}
    if level in font_sizes:
        for run in p.runs:
            run.font.size = font_sizes[level]


def _clear_heading_run(p):
    """清除标题默认 run。"""
    for r in p.runs:
        r.text = ''


def _add_code_block(doc, lines, lang=''):
    """添加代码块（使用等宽字体灰色背景）。"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 灰色背景通过 shading
        shading = run._element.get_or_add_rPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F0F0F0',
            qn('w:val'): 'clear',
        })
        shading.append(shd)


def _add_list(doc, items, ordered=False):
    """添加列表。"""
    for idx, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        prefix = f'{idx+1}.' if ordered else '•'
        run = p.add_run(f'{prefix} ')
        _add_formatted_runs(p, item)


_table_counter = 0  # 表格序号计数器


def _add_table(doc, lines):
    """添加表格（浅蓝表头、100%宽度、居中、自动编号标题）。"""
    global _table_counter
    rows = []
    for line in lines:
        if re.match(r'^\|?[\s\-:|]+\|?$', line.strip()):
            continue  # 分隔行跳过
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if len(rows) < 2:
        return

    _table_counter += 1

    # 表格上方标题
    caption = doc.add_paragraph()
    caption.paragraph_format.line_spacing = 1.5
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f'表格 {_table_counter}')
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'

    # 设置表格宽度为100%
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = tbl.makeelement(qn('w:tblPr'), {})
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = tblPr.makeelement(qn('w:tblW'), {})
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j >= len(table.rows[i].cells):
                continue
            cell = table.rows[i].cells[j]
            cell.text = ''

            # 表头行：浅蓝背景 + 加粗居中
            if i == 0:
                _set_cell_shading(cell, 'B4C6E7')
                run = cell.paragraphs[0].add_run(cell_text)
                run.bold = True
                run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.size = Pt(10)
                # 首列也加粗（通常是指标名）
                if j == 0:
                    run.bold = True
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_cell_shading(cell, color):
    """设置单元格背景色。"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    tcPr.append(shd)


def _add_blockquote(doc, text):
    """添加引用块。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_horizontal_rule(doc):
    """添加水平线。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '999999',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_formatted_runs(p, text):
    """Parse inline formatting and add runs. Handles **bold**, *italic*, `code`, and $formulas$."""
    _PH = '___PROTECT_'
    protected = []

    def _protect(pattern, text):
        parts_list = []
        last = 0
        for m in re.finditer(pattern, text):
            parts_list.append(text[last:m.start()])
            idx = len(protected)
            protected.append(m.group(0))
            parts_list.append(f'{_PH}{idx}___')
            last = m.end()
        parts_list.append(text[last:])
        return ''.join(parts_list)

    # Protect formula and code regions first (they can contain * and other special chars)
    text = _protect(r'\$\$[\s\S]+?\$\$', text)  # display formulas (multiline)
    text = _protect(r'\$[^$\n\r]+\$', text)     # inline formulas
    text = _protect(r'`[^`]+`', text)            # inline code

    # Split by bold/italic markers. Each part may contain regular text + placeholders.
    parts = re.split(r'(\*\*[\s\S]+?\*\*|(?<!\*)\*[^*\s][^*\n]*\*(?!\*))', text)

    for part in parts:
        if not part:
            continue
        # Determine bold/italic wrapper
        bold = False
        italic = False
        content = part
        if part.startswith('**') and part.endswith('**'):
            bold = True
            content = part[2:-2]
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            italic = True
            content = part[1:-1]

        # Within the content, there may be embedded placeholders
        _add_content_with_placeholders(p, content, protected, _PH, bold, italic)


def _add_content_with_placeholders(p, text, protected, _PH, bold=False, italic=False):
    """Add runs from text that may contain ___PROTECT_N___ placeholders."""
    pattern = re.compile(f'{re.escape(_PH)}(\\d+)___')
    last = 0
    for m in pattern.finditer(text):
        # Add text before the placeholder
        if m.start() > last:
            _append_run(p, text[last:m.start()], bold, italic)
        # Add the protected content (formula or code)
        idx = int(m.group(1))
        raw = protected[idx]
        # For inline code, strip the backtick delimiters
        if raw.startswith('`') and raw.endswith('`'):
            run = _append_run(p, raw[1:-1], bold, italic)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            _append_run(p, raw, bold, italic)
        last = m.end()
    # Add remaining text after last placeholder
    if last < len(text):
        _append_run(p, text[last:], bold, italic)


def _append_run(p, text, bold=False, italic=False):
    """Add a run with formatting. Returns the run object."""
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


# ============================================================
# 主入口（测试用）
# ============================================================

if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("用法: python md2docx.py <input.md> [output.docx]")
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(inp)[0] + '.docx'
    md_to_docx(inp, out)
    print(f"Converted: {inp} -> {out}")
