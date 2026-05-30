#!/usr/bin/env python3
"""
convert.py - AI公式 → Word/WPS 专业公式 转换工具

支持三种公式格式:
  1. $...$ / $$...$$ (Markdown风格)
  2. \(...\) / \[...\] (LaTeX风格)
  3. 裸LaTeX (无分隔符，由中英文边界自动识别)

纯 python-docx + lxml, 兼容 WPS 和 Microsoft Word。
"""

import re, os, sys, copy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from latex2omml import latex_to_omml_professional

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ============================================================
# 公式扫描器
# ============================================================

def find_formulas_in_text(text: str):
    """扫描文本中所有公式，返回 (start, end, content, is_display) 列表，从后往前排序。"""
    matches = []
    occupied = []

    def is_free(s, e):
        return not any(o[0] <= s < o[1] or o[0] < e <= o[1] for o in occupied)

    # 第一遍：块级公式
    for pattern, is_disp in [(r'\$\$(.+?)\$\$', True), (r'\\\[(.+?)\\\]', True)]:
        for m in re.finditer(pattern, text, re.DOTALL):
            if not is_free(m.start(), m.end()):
                continue
            if m.start() > 0 and text[m.start() - 1] == '\\':
                continue
            content = m.group(1)
            if not content.strip():
                continue
            matches.append((m.start(), m.end(), content, is_disp))
            occupied.append((m.start(), m.end()))

    # 第二遍：行内公式
    inline_patterns = [
        (r'\$([^\$\n\r]+?)\$', '$', '$'),
        (r'\\\((.*?)\\\)', r'\(', r'\)'),
    ]
    for pattern, d_open, d_close in inline_patterns:
        for m in re.finditer(pattern, text, 0):
            ms, me = m.start(), m.end()
            if not is_free(ms, me):
                continue
            if ms > 0 and text[ms - 1] == '\\':
                continue
            if d_open == '$':
                if ms > 0 and text[ms - 1] == '$':
                    continue
                if me < len(text) and text[me] == '$':
                    continue

            content = m.group(1)
            if not content.strip():
                continue
            if d_open == '$' and _looks_like_currency(content):
                continue

            matches.append((ms, me, content, False))
            occupied.append((ms, me))

    # 第三遍：裸 LaTeX（无分隔符）
    bare = _find_bare_latex(text, occupied)
    matches.extend(bare)
    for bs, be, _, _ in bare:
        occupied.append((bs, be))

    matches.sort(key=lambda x: x[0], reverse=True)
    return matches


def _find_bare_latex(text, occupied):
    """查找无分隔符的 LaTeX 公式。以中文字符为边界分割文本，
    对每个片段检测是否含 LaTeX 命令+数学符号。"""
    results = []

    # 按中文/非中文边界分割
    segments = _split_by_boundary(text)

    for seg_start, seg_end in segments:
        if seg_end - seg_start < 6:
            continue
        if any(o[0] <= seg_start < o[1] or o[0] < seg_end <= o[1] for o in occupied):
            continue

        seg_text = text[seg_start:seg_end]
        if not _has_math_markers(seg_text):
            continue

        # 从片段中提取公式子区域（紧邻 LaTeX 命令的连续数学字符）
        sub_ranges = _extract_math_spans(text, seg_start, seg_end, occupied)
        for s, e in sub_ranges:
            if e - s < 5:
                continue
            content = text[s:e]
            if content.strip():
                results.append((s, e, content, False))
                occupied.append((s, e))

    return results


def _split_by_boundary(text):
    """按非数学边界字符分割文本，返回 (start, end) 列表。
    中文字符、中文标点、换行符等作为边界。"""
    segments = []
    i = 0
    seg_start = 0

    # 边界字符: 中文、中文标点、换行、段落标记
    boundary_set = set()
    # CJK unified ideographs
    for cp in range(0x4E00, 0x9FFF + 1):
        boundary_set.add(chr(cp))
    # CJK compatibility, symbols, punctuation
    for cp in range(0x3000, 0x303F + 1):
        boundary_set.add(chr(cp))
    for cp in range(0xFF00, 0xFFEF + 1):
        boundary_set.add(chr(cp))
    boundary_set.update('，。！？：；""''（）【】《》…—·、×÷≈≠≤≥±∞①②③④⑤⑥⑦⑧⑨⑩→←↑↓↔⇒⇐⇑⇓')
    boundary_set.update('\n\r\t')

    while i < len(text):
        if text[i] in boundary_set:
            if i > seg_start and i - seg_start >= 6:
                segments.append((seg_start, i))
            seg_start = i + 1
        i += 1

    if len(text) - seg_start >= 6:
        segments.append((seg_start, len(text)))

    return segments


def _has_math_markers(seg_text):
    """检查片段是否包含足够的数学标记。"""
    math_indicators = [
        r'\\frac', r'\\dfrac', r'\\sqrt', r'\\sum', r'\\int',
        r'\\prod', r'\\lim', r'\\begin', r'\\end',
        r'\\bar', r'\\vec', r'\\hat', r'\\dot',
        r'\\sin', r'\\cos', r'\\tan',
        r'\\alpha', r'\\beta', r'\\gamma', r'\\delta', r'\\theta',
        r'\\lambda', r'\\mu', r'\\pi', r'\\sigma', r'\\omega',
        r'\\times', r'\\cdot', r'\\pm', r'\\infty',
        r'\\left', r'\\right', r'\\text',
    ]
    cmd_count = sum(1 for ind in math_indicators if ind in seg_text)
    if cmd_count >= 1:
        return True

    # 检测 _ 或 ^ 紧邻字母/数字的模式 (如 v_0, x^2, E_{k1})
    if re.search(r'[a-zA-Z0-9][\^_][a-zA-Z0-9{]', seg_text):
        return True

    # 检测带 {} 的模式 (如 E_{k1})
    if re.search(r'[a-zA-Z][\^_]\\{', seg_text) or re.search(r'\{[^}]*\}', seg_text):
        return True

    return False


def _extract_math_spans(text, seg_start, seg_end, occupied):
    """从片段中提取连续的数学表达式子区域。"""
    math_chars = set(r'^_{}[]()+-*/=<>|!.,;:' + "'\"" + 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789')
    spans = []

    i = seg_start
    while i < seg_end:
        # 找到 LaTeX 命令或 `^` / `_` 作为锚点
        anchor = -1
        if text[i] == '\\':
            anchor = i
        elif text[i] == '^' or text[i] == '_':
            # 往回找最近的字母数字作为公式开头
            j = i - 1
            while j >= seg_start and (text[j].isalnum() or text[j] in math_chars):
                j -= 1
            anchor = j + 1
        elif i > seg_start and text[i] in '{}' and (text[i-1].isalnum() or text[i-1] in '^_'):
            anchor = i

        if anchor >= 0 and not any(o[0] <= anchor < o[1] for o in occupied):
            # 向左扩展到公式边界
            left = anchor
            while left > seg_start and (text[left - 1].isalnum() or text[left - 1] in math_chars or text[left - 1] == '\\'):
                left -= 1
            # 向右扩展
            right = anchor + 1
            while right < seg_end and (text[right].isalnum() or text[right] in math_chars or text[right] == '\\'):
                right += 1

            if right - left >= 5:
                spans.append((left, right))
                i = right
                continue
        i += 1

    # 合并重叠/相邻的 span
    if not spans:
        return []
    spans.sort()
    merged = [spans[0]]
    for s, e in spans[1:]:
        prev_s, prev_e = merged[-1]
        if s <= prev_e + 3:  # 相邻或重叠
            merged[-1] = (prev_s, max(prev_e, e))
        else:
            merged.append((s, e))
    return merged


def _looks_like_currency(content):
    """检测是否像货币金额。"""
    content = content.strip()
    if not content:
        return True
    if content[0].isdigit() or content[0] in ('.', ','):
        for sym in ['\\', '^', '_', '{', '}', '+', '-', '*', '=', '<', '>', '/']:
            if sym in content:
                return False
        return True
    return False


# ============================================================
# 代码块检测
# ============================================================

def _is_inside_code_block(full_text, match_start, match_end):
    """检查匹配是否在代码块/行内代码中。"""
    text_before = full_text[:match_start]
    if text_before.count("```") % 2 == 1:
        return True
    line_start = text_before.rfind('\n') + 1 if '\n' in text_before else 0
    line_before = text_before[line_start:]
    if line_before.count('`') % 2 == 1:
        return True
    return False


# ============================================================
# 文档处理
# ============================================================

def process_paragraph(para):
    """处理单个段落中的公式。返回 (converted, skipped)。"""
    p_elem = para._element
    converted = 0

    for _ in range(60):
        runs_data = _rebuild_runs(p_elem)
        full_text = "".join(r[1] for r in runs_data)
        if not full_text:
            break

        matches = find_formulas_in_text(full_text)
        if not matches:
            break

        ms, me, content, is_display = matches[0]

        if _is_inside_code_block(full_text, ms, me):
            break

        if ms > 0 and full_text[ms - 1] == '\\':
            continue

        try:
            omath_elem = latex_to_omml_professional(content, is_display)
            if _replace_text_with_omath(p_elem, runs_data, ms, me, omath_elem):
                converted += 1
        except Exception:
            pass

    return converted, 0


def _rebuild_runs(p_elem):
    """重建段落运行列表，保留所有非文本 run。"""
    runs_data = []
    for r_elem in p_elem.findall(qn("w:r")):
        if r_elem.find(qn("m:oMath")) is not None or r_elem.find(qn("m:oMathPara")) is not None:
            runs_data.append((r_elem, "", []))
            continue
        if r_elem.find(qn("w:drawing")) is not None or r_elem.find(qn("w:object")) is not None:
            runs_data.append((r_elem, "", []))
            continue
        t_elems = r_elem.findall(qn("w:t"))
        if t_elems:
            text = "".join(t.text or "" for t in t_elems)
            runs_data.append((r_elem, text, t_elems))
        else:
            runs_data.append((r_elem, "", []))
    return runs_data


def _replace_text_with_omath(p_elem, runs_data, ms, me, omath_elem):
    """替换匹配文本为 OMML 公式。保持 前文 → 公式 → 后文 的顺序。"""
    # 找到所在 run
    idx = -1
    for i, (r_elem, text, _) in enumerate(runs_data):
        if not text:
            continue
        run_start = sum(len(r[1]) for r in runs_data[:i])
        if run_start <= ms < run_start + len(text):
            idx = i
            break

    if idx == -1:
        return False

    r_elem, run_text, _ = runs_data[idx]
    run_start = sum(len(r[1]) for r in runs_data[:idx])
    rel_s = ms - run_start
    rel_e = me - run_start

    before_text = run_text[:rel_s]
    after_text = run_text[rel_e:]

    parent = r_elem.getparent()
    r_idx = list(parent).index(r_elem)
    rpr = r_elem.find(qn("w:rPr"))
    rpr_copy = copy.deepcopy(rpr) if rpr is not None else None

    parent.remove(r_elem)

    # 插入: after → omath → before (insert 语义使最终顺序变反，得正向迭代得到正确顺序)
    for text_part in [after_text, None, before_text]:
        if text_part is None:
            nr = etree.SubElement(parent, qn("w:r"), nsmap={"w": WORD_NS})
            if rpr_copy is not None:
                nr.insert(0, copy.deepcopy(rpr_copy))
            nr.insert(1, omath_elem)
        elif text_part:
            nr = etree.SubElement(parent, qn("w:r"), nsmap={"w": WORD_NS})
            if rpr_copy is not None:
                nr.insert(0, copy.deepcopy(rpr_copy))
            nt = etree.SubElement(nr, qn("w:t"))
            nt.text = text_part
        else:
            continue

        children = list(parent)
        if r_idx < len(children):
            parent.insert(r_idx, nr)
        else:
            parent.append(nr)

    return True


# ============================================================
# 主入口
# ============================================================

def process_document(input_path, output_path, auto_number=True):
    """处理整个文档。支持 .docx/.docm 和 .md 文件。"""
    ext = os.path.splitext(input_path)[1].lower()

    # .md → .docx 先转换
    if ext == '.md':
        import tempfile
        tmp_dir = tempfile.mkdtemp()
        docx_tmp = os.path.join(tmp_dir, '_md_temp.docx')
        from md2docx import md_to_docx
        print(f"转换 Markdown → DOCX: {input_path}")
        md_to_docx(input_path, docx_tmp, auto_number=auto_number)
        input_path = docx_tmp
        # 确保输出为 .docx
        if not output_path.lower().endswith('.docx'):
            output_path = os.path.splitext(output_path)[0] + '.docx'

    print(f"打开文档: {input_path}")
    doc = Document(input_path)

    # Process body paragraphs
    converted = sum(process_paragraph(p)[0] for p in doc.paragraphs)

    # Process table cell paragraphs (doc.paragraphs excludes these)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    converted += process_paragraph(para)[0]

    print(f"保存文档: {output_path}")
    doc.save(output_path)
    print(f"\n转换完成！成功转换: {converted} 个公式")
    return converted, 0


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        print("用法: python convert.py <输入文件.docx/.md> [输出文件.docx]")
        sys.exit(1)

    inp = sys.argv[1]
    if not os.path.exists(inp):
        print(f"错误: 文件不存在 - {inp}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        out = sys.argv[2]
    else:
        base, ext = os.path.splitext(inp)
        out_ext = '.docx' if ext.lower() == '.md' else ext
        out = f"{base}_公式转换{out_ext}"

    process_document(inp, out)


if __name__ == "__main__":
    main()
